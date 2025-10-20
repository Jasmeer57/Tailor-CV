"""
File Handler Module
Handles file upload, text extraction, and document creation
"""

import io
from typing import Optional
from docx import Document
import PyPDF2
import pdfplumber


class FileHandler:
    """Handles file operations for CV processing"""

    def __init__(self):
        pass

    def extract_text_from_file(self, uploaded_file) -> Optional[str]:
        """
        Extract text from uploaded PDF or DOCX file

        Args:
            uploaded_file: Streamlit uploaded file object

        Returns:
            Extracted text or None on error
        """
        file_type = uploaded_file.name.split('.')[-1].lower()

        try:
            if file_type == 'pdf':
                return self._extract_from_pdf(uploaded_file)
            elif file_type == 'docx':
                return self._extract_from_docx(uploaded_file)
            else:
                print(f"Unsupported file type: {file_type}")
                return None
        except Exception as e:
            print(f"Error extracting text: {e}")
            return None

    def _extract_from_pdf(self, file) -> str:
        """Extract text from PDF using pdfplumber (more reliable)"""
        try:
            text = ""
            with pdfplumber.open(file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text.strip()
        except:
            # Fallback to PyPDF2
            return self._extract_from_pdf_pypdf2(file)

    def _extract_from_pdf_pypdf2(self, file) -> str:
        """Fallback PDF extraction using PyPDF2"""
        try:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            print(f"PyPDF2 extraction error: {e}")
            return ""

    def _extract_from_docx(self, file) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            print(f"DOCX extraction error: {e}")
            return ""

    def create_docx(self, text: str, title: str = "Document") -> bytes:
        """
        Create a DOCX file from text

        Args:
            text: Text content
            title: Document title

        Returns:
            DOCX file as bytes
        """
        try:
            doc = Document()

            # Add title
            doc.add_heading(title, 0)

            # Add content (split by paragraphs)
            paragraphs = text.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    # Check if it's a heading (simple heuristic)
                    if len(para) < 100 and para.isupper():
                        doc.add_heading(para, level=1)
                    elif len(para) < 100 and ':' in para and para.split(':')[0].isupper():
                        doc.add_heading(para, level=2)
                    else:
                        doc.add_paragraph(para)

            # Save to bytes
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)

            return docx_buffer.getvalue()
        except Exception as e:
            print(f"Error creating DOCX: {e}")
            return b""

    def create_txt(self, text: str) -> bytes:
        """Create a TXT file from text"""
        return text.encode('utf-8')


if __name__ == "__main__":
    # Test file handler
    handler = FileHandler()

    # Test DOCX creation
    sample_text = """John Doe
    Data Scientist

    Experience:
    - 5 years in ML
    - Python expert

    Education:
    - MSc Computer Science"""

    docx_bytes = handler.create_docx(sample_text, "Test CV")
    print(f"Created DOCX: {len(docx_bytes)} bytes")
